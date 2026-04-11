"""
输出格式化与保存（Word + Markdown 双格式）
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class OutputFormatter:

    @staticmethod
    def save_to_file(original_url: str, titles: list, article: str, filepath: str):
        """同时保存 .docx 和 .md 两份文件"""
        path = Path(filepath)
        path.parent.mkdir(parents=True, exist_ok=True)

        # 保存 Word
        OutputFormatter._save_docx(original_url, titles, article, path)

        # 保存 Markdown（同名 .md）
        md_path = path.with_suffix(".md")
        OutputFormatter._save_md(original_url, titles, article, md_path)

    @staticmethod
    def _save_md(original_url: str, titles: list, article: str, path: Path):
        """保存为 Markdown 文件（保留原始格式）"""
        lines = []
        lines.append(f"> 原文链接: {original_url}\n")
        lines.append("## 备选标题\n")
        for i, title in enumerate(titles, 1):
            lines.append(f"{i}. {title}")
        lines.append("\n---\n")
        lines.append("## 正文\n")
        lines.append(article)
        path.write_text("\n".join(lines), encoding="utf-8")

    @staticmethod
    def _save_docx(original_url: str, titles: list, article: str, path: Path):
        """保存为 Word 文档（纯文本）"""
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)

        # 原文链接
        p = doc.add_paragraph()
        run = p.add_run("原文链接: ")
        run.bold = True
        p.add_run(original_url)

        doc.add_paragraph("")

        # 标题列表
        doc.add_heading("备选标题", level=2)
        for i, title in enumerate(titles, 1):
            doc.add_paragraph(f"{i}. {title}")

        doc.add_paragraph("")

        # 正文
        doc.add_heading("正文", level=2)
        for para_text in article.split("\n"):
            text = para_text.strip()
            if text:
                doc.add_paragraph(text)

        doc.save(str(path))
